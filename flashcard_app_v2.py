import streamlit as st 
import ollama
from PyPDF2 import PdfReader
import pandas as pd
from docx import Document
import re

st.title("📚 Simple AI Flashcard Generator")

# File Upload (PDF Only)
uploaded_file = st.file_uploader("📂 Upload your study material (PDF)", type="pdf")

if uploaded_file is not None:
    try:
        # Read PDF
        reader = PdfReader(uploaded_file)

        # Ensure the document has pages
        total_pages = len(reader.pages)
        if total_pages == 0:
            st.error("❌ This PDF has no readable pages. Try another document.")
        else:
            # Input boxes for start and end page
            start_page = st.number_input("📄 Start Page", min_value=1, max_value=total_pages, value=1, step=1)
            end_page = st.number_input("📄 End Page", min_value=start_page, max_value=total_pages, value=total_pages, step=1)

            # Extract text from selected page range
            text = "\n".join([reader.pages[i - 1].extract_text() for i in range(start_page, end_page + 1) if reader.pages[i - 1].extract_text()])

            # If no text is extracted, show an error
            if not text.strip():
                st.error("⚠️ No readable text found in the selected pages. It may be a scanned document or an image-based PDF.")
            else:
                # Display text preview
                st.text_area("📖 Extracted Text Preview", text[:1000])

                # Input box for the number of flashcards
                num_questions = st.number_input("📌 Number of Flashcards", min_value=1, max_value=100, value=10, step=1)

                # Button to generate flashcards
                if st.button("🚀 Generate Flashcards"):
                    with st.spinner("Generating flashcards... please wait ⏳"):
                        response = ollama.chat(
                            model="deepseek-r1",  # Ensure we're using DeepSeek
                            messages=[
                                {
                                    "role": "system",
                                    "content": "You are an AI tutor. Your ONLY task is to generate structured Q&A flashcards. "
                                               "You MUST use the provided document ONLY. "
                                               "DO NOT summarize. DO NOT analyze. DO NOT provide explanations outside the given text."
                                               "Format your response STRICTLY like this:\n\n"
                                               "Q: [Question based on a key concept]\n"
                                               "A: [Answer directly from the text]\n\n"
                                               "Important Rules:\n"
                                               "1. Every Q&A must be based ONLY on the document.\n"
                                               "2. Do NOT provide additional explanations or context.\n"
                                               "3. If a concept is unclear in the text, SKIP it instead of guessing.\n"
                                               "4. Ensure answers are concise but complete.\n\n"
                                               "Example:\n"
                                               "Q: What is photosynthesis?\n"
                                               "A: Photosynthesis is the process where plants convert sunlight into energy.\n\n"
                                               "Now, generate flashcards strictly in this format."
                                },
                                {
                                    "role": "user",
                                    "content": f"Document Name: {uploaded_file.name}\n\n"
                                               f"Generate {num_questions} flashcards in the following format:\n"
                                               f"Q: [Question]\nA: [Answer]\n\n"
                                               f"Each question must be directly based on this text:\n{text}"
                                }
                            ]
                        )

                    try:
                        flashcards_text = response["message"]["content"].strip()  # Ensure valid extraction
                    except KeyError:
                        st.error("❌ AI did not return a valid response. Try again.")
                        flashcards_text = ""

                    if not flashcards_text:
                        st.error("❌ No flashcards generated. Try again with different settings.")
                    else:
                        # Improved regex for extracting Q&A pairs
                        flashcards = re.findall(r"Q:\s*(.+?)\s*\nA:\s*(.+)", flashcards_text, re.DOTALL)

                        # Convert to DataFrame with stripped spaces
                        df = pd.DataFrame([{"Question": q.strip(), "Answer": a.strip()} for q, a in flashcards])
                        
                        # Prevent empty downloads
                        if df.empty:
                            st.error("❌ No valid questions found in AI output.")
                        else:
                            st.subheader("✅ Generated Flashcards:")
                            st.write(df)

                            # Generate filenames
                            base_filename = uploaded_file.name.split('.')[0]
                            csv_filename = f"{base_filename}_flashcards.csv"
                            doc_filename = f"{base_filename}_flashcards.docx"

                            # Export as CSV
                            csv = df.to_csv(index=False).encode('utf-8')
                            st.download_button("📥 Download Flashcards as CSV", csv, csv_filename, "text/csv")

                            # Export as DOCX
                            doc = Document()
                            doc.add_heading(f"Generated Flashcards from {uploaded_file.name}", level=1)

                            for _, row in df.iterrows():
                                doc.add_paragraph(row["Question"])
                                doc.add_paragraph(row["Answer"])
                                doc.add_paragraph("\n")  # Spacing for clarity

                            doc.save(doc_filename)
                            with open(doc_filename, "rb") as f:
                                st.download_button("📥 Download Flashcards as DOCX", f, doc_filename,
                                                    "application/vnd.openxmlformats-officedocument.wordprocessingml.document")

    except Exception as e:
        st.error(f"❌ An error occurred: {e}")
